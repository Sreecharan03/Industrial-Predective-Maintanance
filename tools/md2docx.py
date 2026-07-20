"""Markdown -> Word converter for the SenseMinds 360 docs.

Handles the constructs the docs actually use: headings, pipe tables, fenced code
blocks, bullet/numbered/checkbox lists, blockquotes, rules, and inline
bold / `code` / [links].
"""

from __future__ import annotations

import pathlib
import re
import sys

from docx import Document
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor

INK = RGBColor(0x1C, 0x19, 0x17)
MUTED = RGBColor(0x57, 0x53, 0x4E)
BRAND = RGBColor(0x7C, 0x3A, 0xED)
CODE_INK = RGBColor(0xBE, 0x12, 0x3C)
CODE_BLOCK_INK = RGBColor(0x1C, 0x19, 0x17)
LINK = RGBColor(0x02, 0x84, 0xC7)


def shade(el, hex_colour):
    node = OxmlElement("w:shd")
    node.set(qn("w:fill"), hex_colour)
    el.get_or_add_tcPr().append(node) if hasattr(el, "get_or_add_tcPr") else None


def shade_cell(cell, hex_colour):
    node = OxmlElement("w:shd")
    node.set(qn("w:fill"), hex_colour)
    cell._tc.get_or_add_tcPr().append(node)


def shade_para(paragraph, hex_colour):
    node = OxmlElement("w:shd")
    node.set(qn("w:val"), "clear")
    node.set(qn("w:fill"), hex_colour)
    paragraph._p.get_or_add_pPr().append(node)


_INLINE = re.compile(r"(\*\*.+?\*\*|`[^`]+`|\[[^\]]+\]\([^)]+\))")


def split_cells(row: str) -> list[str]:
    """Split a pipe-table row, honouring escaped '\\|' inside a cell."""
    parts = re.split(r"(?<!\\)\|", row.strip().strip("|"))
    return [p.strip().replace("\\|", "|") for p in parts]


def add_inline(paragraph, text, *, base_size=10.5, bold=False, color=None):
    """Render markdown inline formatting into runs."""
    text = text.replace("&middot;", "·").replace("&nbsp;", " ").replace("\\|", "|")
    for part in _INLINE.split(text):
        if not part:
            continue
        if part.startswith("**") and part.endswith("**"):
            r = paragraph.add_run(part[2:-2])
            r.bold = True
        elif part.startswith("`") and part.endswith("`"):
            r = paragraph.add_run(part[1:-1])
            r.font.name = "Consolas"
            r.font.color.rgb = CODE_INK
            r.font.size = Pt(base_size - 0.5)
        elif part.startswith("[") and "](" in part:
            label = part[1 : part.index("]")]
            r = paragraph.add_run(label)
            r.font.color.rgb = LINK
            r.underline = True
        else:
            r = paragraph.add_run(part)
            r.bold = bold
        r.font.size = Pt(base_size)
        if color is not None and not (part.startswith("`")):
            r.font.color.rgb = color


def _consume_item(lines, i):
    """Join a list item's wrapped continuation lines into one logical line.

    Without this, a bold span or link that happens to wrap across two source
    lines is split into two paragraphs and renders as literal ** markers."""
    parts = [lines[i].strip()]
    i += 1
    while i < len(lines):
        nxt = lines[i]
        stripped = nxt.strip()
        if (not stripped or stripped.startswith(("#", "|", "```", ">", "-", "*", "!["))
                or re.match(r"^\d+\.\s", stripped) or stripped in ("---", "***")):
            break
        if not nxt.startswith((" ", "\t")) and len(parts) == 1:
            # unindented wrap is still part of the item
            pass
        parts.append(stripped)
        i += 1
    return " ".join(parts), i


def build(md_path: str, out_path: str, title: str, subtitle: str) -> None:
    doc = Document()

    normal = doc.styles["Normal"]
    normal.font.name = "Calibri"
    normal.font.size = Pt(10.5)
    normal.font.color.rgb = INK
    normal.paragraph_format.space_after = Pt(7)
    normal.paragraph_format.line_spacing = 1.13

    for name, size, colour, before in (
        ("Heading 1", 16, BRAND, 20),
        ("Heading 2", 12.5, INK, 13),
        ("Heading 3", 11, MUTED, 10),
        ("Heading 4", 10.5, MUTED, 8),
    ):
        st = doc.styles[name]
        st.font.name = "Calibri"
        st.font.size = Pt(size)
        st.font.color.rgb = colour
        st.font.bold = True
        st.paragraph_format.space_before = Pt(before)
        st.paragraph_format.space_after = Pt(5)

    # ---- cover ----
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run("SenseMinds 360")
    r.bold = True
    r.font.size = Pt(24)
    r.font.color.rgb = BRAND

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run(title)
    r.font.size = Pt(14)
    r.bold = True

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run(subtitle)
    r.font.size = Pt(9.5)
    r.font.color.rgb = MUTED
    doc.add_paragraph()

    lines = open(md_path, encoding="utf-8").read().splitlines()
    i = 0
    first_heading_seen = False

    while i < len(lines):
        line = lines[i]
        stripped = line.strip()

        # fenced code block
        if stripped.startswith("```"):
            i += 1
            code: list[str] = []
            while i < len(lines) and not lines[i].strip().startswith("```"):
                code.append(lines[i])
                i += 1
            i += 1
            t = doc.add_table(rows=1, cols=1)
            cell = t.rows[0].cells[0]
            cell.text = ""
            shade_cell(cell, "F5F5F4")
            for n, cl in enumerate(code):
                para = cell.paragraphs[0] if n == 0 else cell.add_paragraph()
                para.paragraph_format.space_after = Pt(0)
                para.paragraph_format.line_spacing = 1.0
                r = para.add_run(cl)
                r.font.name = "Consolas"
                r.font.size = Pt(8.5)
                r.font.color.rgb = CODE_BLOCK_INK
            doc.add_paragraph()
            continue

        # pipe table
        if stripped.startswith("|") and i + 1 < len(lines) and re.match(
            r"^\|[\s:|-]+\|$", lines[i + 1].strip()
        ):
            headers = split_cells(stripped)
            i += 2
            rows = []
            while i < len(lines) and lines[i].strip().startswith("|"):
                rows.append(split_cells(lines[i]))
                i += 1
            t = doc.add_table(rows=1, cols=len(headers))
            t.style = "Table Grid"
            t.alignment = WD_TABLE_ALIGNMENT.CENTER
            for n, h in enumerate(headers):
                c = t.rows[0].cells[n]
                c.text = ""
                add_inline(c.paragraphs[0], h, base_size=9, bold=True)
                for run in c.paragraphs[0].runs:
                    run.bold = True
                shade_cell(c, "EDE9FE")
            for row in rows:
                cells = t.add_row().cells
                for n, val in enumerate(row[: len(headers)]):
                    cells[n].text = ""
                    add_inline(cells[n].paragraphs[0], val, base_size=9)
            width = 6.7 / max(len(headers), 1)
            for r_ in t.rows:
                for c in r_.cells:
                    c.width = Inches(width)
            doc.add_paragraph()
            continue

        # image:  ![caption](path)
        m = re.match(r"^!\[([^\]]*)\]\(([^)]+)\)\s*$", stripped)
        if m:
            caption, src = m.group(1), m.group(2)
            path = (pathlib.Path(md_path).parent / src).resolve()
            if path.exists():
                doc.add_picture(str(path), width=Inches(6.3))
                doc.paragraphs[-1].alignment = WD_ALIGN_PARAGRAPH.CENTER
                if caption:
                    cap = doc.add_paragraph()
                    cap.alignment = WD_ALIGN_PARAGRAPH.CENTER
                    r = cap.add_run(caption)
                    r.italic = True
                    r.font.size = Pt(9)
                    r.font.color.rgb = MUTED
            i += 1
            continue

        # headings
        m = re.match(r"^(#{1,4})\s+(.*)$", stripped)
        if m:
            level = len(m.group(1))
            text = m.group(2)
            if level == 1 and not first_heading_seen:
                first_heading_seen = True
                if text.lower().startswith("senseminds"):
                    i += 1
                    continue
            h = doc.add_heading("", level=min(level, 4))
            add_inline(h, text, base_size=16 - (level - 1) * 2.5)
            for run in h.runs:
                run.bold = True
                run.font.color.rgb = BRAND if level == 1 else (INK if level == 2 else MUTED)
            i += 1
            continue

        # horizontal rule -> spacer
        if stripped in ("---", "***", "___"):
            doc.add_paragraph()
            i += 1
            continue

        # blockquote (possibly multi-line)
        if stripped.startswith(">"):
            quote: list[str] = []
            while i < len(lines) and lines[i].strip().startswith(">"):
                quote.append(lines[i].strip().lstrip(">").strip())
                i += 1
            t = doc.add_table(rows=1, cols=1)
            cell = t.rows[0].cells[0]
            cell.text = ""
            shade_cell(cell, "FEF3C7")
            add_inline(cell.paragraphs[0], " ".join(q for q in quote if q), base_size=10)
            doc.add_paragraph()
            continue

        # checkbox list
        m = re.match(r"^-\s+\[([ xX])\]\s+(.*)$", stripped)
        if m:
            text, i = _consume_item(lines, i)
            body = re.match(r"^-\s+\[[ xX]\]\s+(.*)$", text).group(1)
            p = doc.add_paragraph(style="List Bullet")
            p.add_run("☒ " if m.group(1).lower() == "x" else "☐ ").font.size = Pt(11)
            add_inline(p, body)
            continue

        # bullet
        if re.match(r"^[-*]\s+", stripped):
            text, i = _consume_item(lines, i)
            p = doc.add_paragraph(style="List Bullet")
            add_inline(p, re.sub(r"^[-*]\s+", "", text))
            continue

        # numbered
        if re.match(r"^\d+\.\s+", stripped):
            text, i = _consume_item(lines, i)
            p = doc.add_paragraph(style="List Number")
            add_inline(p, re.sub(r"^\d+\.\s+", "", text))
            continue

        # blank
        if not stripped:
            i += 1
            continue

        # paragraph (join continuation lines)
        para_lines = [stripped]
        i += 1
        while i < len(lines):
            nxt = lines[i].strip()
            if (not nxt or nxt.startswith(("#", "|", "```", ">", "-", "*"))
                    or re.match(r"^\d+\.\s", nxt) or nxt in ("---", "***")):
                break
            para_lines.append(nxt)
            i += 1
        p = doc.add_paragraph()
        add_inline(p, " ".join(para_lines))

    doc.save(out_path)
    print(f"wrote {out_path}")


if __name__ == "__main__":
    build(*sys.argv[1:5])
